from flask import Flask, render_template, request, jsonify, send_file
from flask import session, redirect, url_for, send_from_directory
from functools import wraps
import requests
import os
import random
import hmac
import hashlib
from datetime import datetime, timedelta
from tavily import TavilyClient
from werkzeug.utils import secure_filename
from PIL import Image
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.pdfgen import canvas
from docx import Document
import fitz
import razorpay
from flask_sqlalchemy import SQLAlchemy
import requests

def send_otp_email(to_email, otp):
    url = "https://api.resend.com/emails"

    headers = {
        "Authorization": "Bearer re_ieNDTuCP_8ewyF3N8n9zzUwr4JzDB231j",
        "Content-Type": "application/json"
    }
    html = f"""
    <div style="background:linear-gradient(135deg,#0b0b14,#1a0f2e);
                padding:40px 0;
                font-family:-apple-system,BlinkMacSystemFont,'Segoe UI',Roboto,Arial,sans-serif;">

    <div style="max-width:480px;
                margin:auto;
                background:#12121c;
                padding:32px;
                border-radius:12px;
                border:1px solid rgba(255,255,255,0.08);">

                 <img src="https://www.qerrastar.online/static/images/logo.png"
         alt="Qerrastar"
         style="width:120px;margin-bottom:20px;">


        <h2 style="margin:0 0 20px 0;
                font-size:18px;
                color:#ffffff;
                letter-spacing:1px;">
        Qerrastar
        </h2>

        <p style="font-size:16px;color:#e5e7eb;margin-bottom:20px;">
        Verify your email address
        </p>

        <p style="font-size:14px;color:#a1a1aa;margin-bottom:25px;">
        Enter the verification code below to continue:
        </p>

        <div style="font-size:30px;
                    font-weight:600;
                    letter-spacing:6px;
                    color:#a855f7;
                    margin-bottom:20px;">
        {otp}
        </div>

        <p style="font-size:13px;color:#71717a;margin-bottom:30px;">
        This code will expire in 5 minutes.
        </p>

        <hr style="border:none;border-top:1px solid rgba(255,255,255,0.08);margin:20px 0;">

        <p style="font-size:12px;color:#6b7280;">
        If you didn’t request this, you can safely ignore this email.
        </p>

    </div>

    </div>
    """

    data = {
        "from": "Qerrastar <no-reply@qerrastar.online>",
        "to": [to_email],
        "subject": "Verify Your Qerrastar Account",
        "html": html
    }

    response = requests.post(url, json=data, headers=headers)

    print("STATUS:", response.status_code)
    print("RESPONSE:", response.text)

#For Gmail Contact
def send_contact_email(name, email, subject, body):
    url = "https://api.resend.com/emails"

    headers = {
        "Authorization": "Bearer re_ieNDTuCP_8ewyF3N8n9zzUwr4JzDB231j",
        "Content-Type": "application/json"
    }

    html_content = f"""
    <h3>New Contact Message</h3>
    <p><b>Name:</b> {name}</p>
    <p><b>Email:</b> {email}</p>
    <p><b>Message:</b><br>{body}</p>
    """

    data = {
        "from": "Qerrastar <no-reply@qerrastar.online>",
        "to": ["saikatmahara7895@gmail.com"],
        "subject": f"Qerrastar Contact: {subject}",
        "html": html_content
    }

    response = requests.post(url, json=data, headers=headers)

    print("CONTACT STATUS:", response.status_code)
    print("CONTACT RESPONSE:", response.text)




app = Flask(__name__)
app.secret_key = os.environ.get("SECRET_KEY", "change-this-later")

# ── DATABASE CONFIG ──
# Use DATABASE_URL env var for production (PostgreSQL on Supabase/Neon/etc.)
# Falls back to SQLite only for local development
database_url = os.environ.get("DATABASE_URL", "sqlite:///users.db")
# Some providers give postgres:// instead of postgresql:// — fix that
if database_url.startswith("postgres://"):
    database_url = database_url.replace("postgres://", "postgresql://", 1)

app.config['SQLALCHEMY_DATABASE_URI'] = database_url
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False


from models import db, User

db.init_app(app)

with app.app_context():
    db.create_all()
    # Safely add new columns if they don't exist yet
    try:
        with db.engine.connect() as conn:
            conn.execute(db.text('ALTER TABLE "user" ADD COLUMN is_pro BOOLEAN DEFAULT 0'))
            conn.commit()
    except: pass
    try:
        with db.engine.connect() as conn:
            conn.execute(db.text('ALTER TABLE "user" ADD COLUMN pro_since DATETIME'))
            conn.commit()
    except: pass
    try:
        with db.engine.connect() as conn:
            conn.execute(db.text('ALTER TABLE "user" ADD COLUMN pro_expires DATETIME'))
            conn.commit()
    except: pass
#----------------------------------------------------------------
# FOR SITEMAP
@app.route('/sitemap.xml')
def sitemap():
    return send_from_directory(os.getcwd(), 'sitemap.xml')
#----------------------------------------------------------------

# ---------------- LOGIN REQUIRED DECORATOR ----------------
def login_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if "user" not in session:
            return redirect(url_for("login"))
        return f(*args, **kwargs)
    return decorated_function


# ---------------- ROUTES ----------------

@app.route("/")
def index():
    return render_template("index.html")


# LEARN / SEARCH

TAVILY_API_KEY = os.getenv("TAVILY_API_KEY")
SEARCH_ENGINE_ID = os.getenv("SEARCH_ENGINE_ID")
@app.route("/search")
def search_page():
    return render_template("search.html")

@app.route("/api/search")
def api_search():
    query = request.args.get("q")

    url = "https://api.tavily.com/search"

    payload = {
        "api_key": TAVILY_API_KEY,
        "query": query,
        "search_depth": "advanced",
        "include_answer": False,
        "max_results": 10
    }

    response = requests.post(url, json=payload)
    data = response.json()

    results = []

    for item in data.get("results", []):
        content = item.get("content", "")
        short_snippet = " ".join(content.split()[:7]) + "..."

        results.append({
            "title": item.get("title"),
            "snippet": short_snippet,
            "link": item.get("url")
        })

    return jsonify({"items": results})



# ======================= QUIZ =======================

# Temporary sample question generator
tavily = TavilyClient(api_key=os.getenv("TAVILY_API_KEY"))
def generate_quiz(topic):

    response = tavily.search(
        query=f"{topic} facts explanation definition",
        search_depth="advanced",
        max_results=5
    )

    content = ""
    for result in response["results"]:
        content += result["content"] + " "

    sentences = content.split(".")
    sentences = [s.strip() for s in sentences if len(s.strip()) > 50]

    if len(sentences) < 5:
        return []

    selected = random.sample(sentences, 5)

    questions = []

    for sentence in selected:

        question = f"What concept is described below?\n\n{sentence}"

        # generate random wrong answers
        distractors = [
            f"{topic} theory",
            f"{topic} process",
            f"{topic} principle",
            f"{topic} method",
            f"{topic} system"
        ]

        options = [topic] + random.sample(distractors, 3)

        random.shuffle(options)

        questions.append({
            "question": question,
            "options": options,
            "answer": topic
        })

    return questions


@app.route("/quiz", methods=["GET", "POST"])
def quiz():
    if request.method == "POST":
        topic = request.form["topic"]
        questions = generate_quiz(topic)
        return render_template("quiz.html", questions=questions, topic=topic)
    return render_template("quiz.html", questions=None)


@app.route("/submit_quiz", methods=["POST"])
def submit_quiz():
    score = 0
    total = int(request.form["total"])

    for i in range(total):
        selected = request.form.get(f"q{i}")
        correct = request.form.get(f"correct{i}")
        if selected == correct:
            score += 1

    return render_template("result.html", score=score, total=total)





# ================= SUMMARIZE =================

@app.route("/summary", methods=["GET", "POST"])
def summarize():
    summary = None

    if request.method == "POST":
        try:
            text = request.form.get("topic")

            if not text:
                summary = "Please enter text."
                return render_template("summarize.html", summary=summary)

            # If short input → use Tavily search
            if len(text) <= 400:
                response = tavily.search(
                    query=text,
                    search_depth="basic",
                    max_results=3
                )

                content = ""
                if response and response.get("results"):
                    for result in response["results"]:
                        content += result.get("content", "") + " "
                text = content
            import re

            # Markdown links
            text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
            text = re.sub(r'http\S+', '', text)
            text = re.sub(r'\s+', ' ', text)

            sentences = re.split(r'\.\s+', text)

            cleaned = []
            seen = set()

            for s in sentences:
                s = s.strip()
                if len(s) > 60:
                    key = s.lower()
                    if key not in seen:
                        seen.add(key)
                        cleaned.append(s)

            if not cleaned:
                summary = "Not enough meaningful content to summarize."
            else:
                summary = ". ".join(cleaned[:3]) + "."

        except Exception as e:
            summary = f"Error occurred: {str(e)}"

    return render_template("summarize.html", summary=summary)





# ======================= DASHBOARD =======================
@app.route("/dashboard")
def dashboard():
    return render_template("dashboard.html")


# ======================= DOCUMENTATION =======================
@app.route("/docs")
def docs():
    return render_template("docs.html")


# ======================= CONVERTER =======================
UPLOAD_FOLDER = "uploads"
CONVERTED_FOLDER = "converted"

os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(CONVERTED_FOLDER, exist_ok=True)

app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER
app.config["CONVERTED_FOLDER"] = CONVERTED_FOLDER
# ================= CONVERTER PAGE =================
@app.route("/converter")
def converter():
    return render_template("converter.html")


# ================= JPG → PDF =================
@app.route("/convert/jpg-to-pdf", methods=["POST"])
def jpg_to_pdf():
    file = request.files["file"]
    filename = secure_filename(file.filename)
    filepath = os.path.join(app.config["UPLOAD_FOLDER"], filename)
    file.save(filepath)

    image = Image.open(filepath).convert("RGB")
    output_path = os.path.join(app.config["CONVERTED_FOLDER"], filename.rsplit(".", 1)[0] + ".pdf")
    image.save(output_path)

    return send_file(output_path, as_attachment=True)


# ================= PDF → JPG =================
@app.route("/convert/pdf-to-jpg", methods=["POST"])
def pdf_to_jpg():
    file = request.files["file"]
    filename = secure_filename(file.filename)
    filepath = os.path.join(app.config["UPLOAD_FOLDER"], filename)
    file.save(filepath)

    doc = fitz.open(filepath)
    page = doc[0]
    pix = page.get_pixmap()

    output_path = os.path.join(app.config["CONVERTED_FOLDER"], filename.rsplit(".", 1)[0] + ".jpg")
    pix.save(output_path)

    return send_file(output_path, as_attachment=True)


# ================= WORD → PDF =================
@app.route("/convert/word-to-pdf", methods=["POST"])
def word_to_pdf():
    file = request.files["file"]
    filename = secure_filename(file.filename)
    filepath = os.path.join(app.config["UPLOAD_FOLDER"], filename)
    file.save(filepath)

    doc = Document(filepath)
    text = "\n".join([para.text for para in doc.paragraphs])

    output_path = os.path.join(app.config["CONVERTED_FOLDER"], filename.rsplit(".", 1)[0] + ".pdf")
    pdf = SimpleDocTemplate(output_path)
    styles = getSampleStyleSheet()
    elements = [Paragraph(text, styles["Normal"])]
    pdf.build(elements)

    return send_file(output_path, as_attachment=True)


# ================= PDF → WORD =================
@app.route("/convert/pdf-to-word", methods=["POST"])
def pdf_to_word():
    file = request.files["file"]
    filename = secure_filename(file.filename)
    filepath = os.path.join(app.config["UPLOAD_FOLDER"], filename)
    file.save(filepath)

    doc = fitz.open(filepath)
    text = ""
    for page in doc:
        text += page.get_text()

    output_path = os.path.join(app.config["CONVERTED_FOLDER"], filename.rsplit(".", 1)[0] + ".docx")

    document = Document()
    document.add_paragraph(text)
    document.save(output_path)

    return send_file(output_path, as_attachment=True)

# ================= PYQ =================
BASE_DIR = os.getcwd()
PYQ_FOLDER = os.path.join(BASE_DIR, "pyq_storage")

os.makedirs(PYQ_FOLDER, exist_ok=True)

app.config["PYQ_FOLDER"] = PYQ_FOLDER
@app.route("/pyq")  #ROUTE
def pyq():
    return render_template("pyq.html")
app.config['MAX_CONTENT_LENGTH'] = 3 * 1024 * 1024  # 3MB limit
#++++++++++++++ Search & Download ++++++++++++++++
@app.route("/pyq/search", methods=["POST"])
def search_pyq():

    university = secure_filename(request.form.get("university")).lower()
    degree = secure_filename(request.form.get("degree")).lower()
    department = secure_filename(request.form.get("department")).lower()
    year = secure_filename(request.form.get("year")).lower()
    subject = secure_filename(request.form.get("subject")).lower()

    file_path = os.path.join(
        app.config["PYQ_FOLDER"],
        university,
        degree,
        department,
        year,
        f"{subject}.pdf"
    )

    if os.path.exists(file_path):
        return send_file(file_path, as_attachment=True)
    else:
        return "PYQ not found."
    
#++++++++++++++ UPLOAD ROUTE +++++++++++++
@app.route("/pyq/upload", methods=["POST"])
def upload_pyq():
    university_input = request.form.get("university")
    degree = request.form.get("degree")
    department = request.form.get("department")
    year = request.form.get("year")
    subject = request.form.get("subject")
    file = request.files.get("file")

    if not all([university_input, degree, department, year, subject, file]):
        return "Missing required fields"

    base_root = app.config["PYQ_FOLDER"]

    # 🔥 Case-insensitive match for existing university
    matched_university = None
    for folder in os.listdir(base_root):
        if folder.lower() == university_input.lower():
            matched_university = folder
            break

    if not matched_university:
        # If not exists → create new
        matched_university = secure_filename(university_input)

    # Clean inputs
    degree = secure_filename(degree)
    department = secure_filename(department)
    year = secure_filename(year)
    subject = secure_filename(subject)

    folder_path = os.path.join(
        base_root,
        matched_university,
        degree,
        department,
        year
    )

    os.makedirs(folder_path, exist_ok=True)

    # 🔥 Prevent overwrite (auto-numbering)
    file_path = os.path.join(folder_path, f"{subject}.pdf")
    counter = 1

    while os.path.exists(file_path):
        file_path = os.path.join(folder_path, f"{subject}_{counter}.pdf")
        counter += 1

    file.save(file_path)

    return "Uploaded Successfully"
    
#======================= LIST ======================
@app.route("/pyq/list", methods=["POST"])
def list_pyq():
    university_input = request.form.get("university")

    if not university_input:
        return "University required"

    base_root = app.config["PYQ_FOLDER"]

    # 🔥 Case-insensitive match for university folder
    matched_university = None
    for folder in os.listdir(base_root):
        if folder.lower() == university_input.lower():
            matched_university = folder
            break

    if not matched_university:
        return "No records found"

    base_path = os.path.join(base_root, matched_university)

    files = []

    # 🔥 Walk through ALL subfolders
    for root, dirs, filenames in os.walk(base_path):
        for filename in filenames:
            if filename.endswith(".pdf"):
                relative_path = os.path.relpath(
                    os.path.join(root, filename),
                    base_path
                ).replace("\\", "/")

                files.append(relative_path)

    return render_template(
        "pyq_list.html",
        files=files,
        university=matched_university
    )
#======================= DOWNLOAD ROUTE ===================
@app.route("/pyq/download/<path:filepath>")
def download_pyq(filepath):
    university = request.args.get("university")

    full_path = os.path.join(
        app.config["PYQ_FOLDER"],
        university,
        filepath
    )

    if os.path.exists(full_path):
        return send_file(full_path, as_attachment=True)

    return "File not found"

# ============================================================
# L O G I N  &  S I G N  U P
# ============================================================

# ── LOGIN ──
@app.route("/login", methods=["GET", "POST"])
def login():
    error = None
    if request.method == "POST":
        email    = request.form.get("email", "").strip()
        password = request.form.get("password", "")

        user = User.query.filter_by(email=email).first()

        if not user:
            error = "No account found with that email."
        elif not user.check_password(password):
            error = "Incorrect password. Please try again."
        elif not user.verified:
            error = "Please verify your email first. Check your inbox."
        else:
            session["user"] = user.email
            return redirect(url_for("dashboard"))

    return render_template("login.html", error=error)


# ── SIGNUP ──
@app.route("/signup", methods=["GET", "POST"])
def signup():
    error = None
    if request.method == "POST":
        email    = request.form.get("email", "").strip()
        password = request.form.get("password", "")

        if len(password) < 6:
            error = "Password must be at least 6 characters."
            return render_template("signup.html", error=error)

        existing_user = User.query.filter_by(email=email).first()
        if existing_user:
            error = "This email is already registered. Try logging in."
            return render_template("signup.html", error=error)

        otp      = str(random.randint(100000, 999999))
        new_user = User(email=email, otp=otp)
        new_user.set_password(password)
        db.session.add(new_user)
        db.session.commit()

        try:
            send_otp_email(email, otp)
        except Exception as e:
            print("EMAIL ERROR:", e)
            return "Failed to send OTP"
        return redirect(url_for("verify", email=email))

    return render_template("signup.html", error=error)


# ── VERIFY OTP ──
@app.route("/verify/<email>", methods=["GET", "POST"])
def verify(email):
    user  = User.query.filter_by(email=email).first()
    error = None

    if request.method == "POST":
        entered_otp = request.form.get("otp", "").strip()

        if not user:
            error = "Account not found."
        elif user.otp != entered_otp:
            error = "Incorrect OTP. Please check your email and try again."
        else:
            user.verified = True
            user.otp      = None
            db.session.commit()
            return redirect(url_for("login"))

    return render_template("verify.html", email=email, error=error)


#LOGOUT ROUTE
@app.route("/logout")
def logout():
    session.clear()
    return redirect(url_for("index"))


# ============================================================
# PREMIUM PLANS — RAZORPAY
# ============================================================

RAZORPAY_KEY_ID     = os.environ.get("RAZORPAY_KEY_ID", "rzp_test_XXXXXXXXXXXXXXXX")
RAZORPAY_KEY_SECRET = os.environ.get("RAZORPAY_KEY_SECRET", "XXXXXXXXXXXXXXXXXXXXXXXX")

rzp_client = razorpay.Client(auth=(RAZORPAY_KEY_ID, RAZORPAY_KEY_SECRET))


@app.route("/pricing")
def pricing():
    return render_template("pricing.html")


@app.route("/create-order", methods=["POST"])
def create_order():
    if "user" not in session:
        return jsonify({"error": "not logged in"}), 401

    data   = request.get_json()
    amount = data.get("amount", 4900)
    period = data.get("period", "monthly")

    order = rzp_client.order.create({
        "amount":   amount,
        "currency": "INR",
        "payment_capture": 1
    })

    return jsonify({
        "key":      RAZORPAY_KEY_ID,
        "amount":   order["amount"],
        "order_id": order["id"]
    })


@app.route("/verify-payment", methods=["POST"])
def verify_payment():
    if "user" not in session:
        return jsonify({"success": False}), 401

    data = request.get_json()

    body     = data["razorpay_order_id"] + "|" + data["razorpay_payment_id"]
    expected = hmac.new(
        key=RAZORPAY_KEY_SECRET.encode(),
        msg=body.encode(),
        digestmod=hashlib.sha256
    ).hexdigest()
    actual   = data["razorpay_signature"]

    if expected != actual:
        return jsonify({"success": False, "error": "Invalid signature"})

    user = User.query.filter_by(email=session["user"]).first()
    if user:
        user.is_pro      = True
        user.pro_since   = datetime.utcnow()
        user.pro_expires = datetime.utcnow() + timedelta(days=30)
        db.session.commit()

    return jsonify({"success": True})


@app.route("/payment-success")
def payment_success():
    return render_template("payment_success.html")

# Privacy terms Routes

@app.route("/about")
def about():
    return render_template("about.html")

@app.route("/contact", methods=["GET", "POST"])
def contact():
    success = None
    error = None
    if request.method == "POST":
        name    = request.form.get("name", "").strip()
        email   = request.form.get("email", "").strip()
        subject = request.form.get("subject", "").strip()
        body    = request.form.get("body", "").strip()

    try:
        send_contact_email(name, email, subject, body)
        success = "Message sent! I'll get back to you within 24–48 hours."
    except Exception as e:
         print("CONTACT ERROR:", e)
         error = "Failed to send. Please try again later."

    return render_template("contact.html", success=success, error=error)

@app.route("/privacy")
def privacy():
    return render_template("privacy.html")

@app.route("/terms")
def terms():
    return render_template("terms.html")

@app.route("/disclaimer")
def disclaimer():
    return render_template("disclaimer.html")


if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=port)